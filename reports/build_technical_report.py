from __future__ import annotations

import csv
import statistics
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(r"C:\Users\asus\Duckcurve_flat_improved")
RESULTS = ROOT / "results_81pct"
REPORTS = ROOT / "reports"
OUTPUT = REPORTS / "Duckcurve_Optimization_Technical_Report.docx"
SKILL_SCRIPTS = Path(
    r"C:\Users\asus\.codex\plugins\cache\openai-primary-runtime"
    r"\documents\26.723.12215\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5B6573"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "FFF5D6"
GOLD = "7A5A00"
RED = "9B1C1C"
GREEN = "1F6B45"
WHITE = "FFFFFF"
BLACK = "111111"


def read_csv(name: str) -> list[dict[str, str]]:
    with (RESULTS / name).open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


summary_rows = read_csv("results_summary.csv")
audit_rows = read_csv("bess_constraint_audit.csv")
resilience_rows = read_csv("resilience_indices.csv")
scenario_rows = read_csv("resilience_scenarios.csv")
hourly_rows = read_csv("results_hourly.csv")


def row_by(rows: list[dict[str, str]], key: str, value: str) -> dict[str, str]:
    return next(row for row in rows if row[key] == value)


baseline = row_by(summary_rows, "label", "duck baseline (PV,no BESS)")
optimized = row_by(summary_rows, "label", "best_duck")
improvement = row_by(summary_rows, "label", "improvement_%")
res_before = row_by(resilience_rows, "scenario", "before")
res_after = row_by(resilience_rows, "scenario", "after")


def f(row: dict[str, str], key: str) -> float:
    return float(row[key])


def pct_lower(before: float, after: float) -> float:
    return (before - after) / before * 100.0


def pct_higher(before: float, after: float) -> float:
    return (after - before) / before * 100.0


scenario_stats = {}
for label in ("before", "after"):
    group = [row for row in scenario_rows if row["scenario"] == label]
    ens = [float(row["unserved_energy_kwh"]) for row in group]
    served = [float(row["load_served_percent"]) for row in group]
    scenario_stats[label] = {
        "ens_mean": statistics.mean(ens),
        "ens_median": statistics.median(ens),
        "served_mean": statistics.mean(served),
        "zero_ens": sum(value == 0.0 for value in ens),
    }

baseline_v = [float(row["v_min_base"]) for row in hourly_rows]
optimized_v = [float(row["v_min_opt"]) for row in hourly_rows]
worst_base_v = min(baseline_v)
worst_opt_v = min(optimized_v)
worst_opt_hour = optimized_v.index(worst_opt_v)
base_hours_low = sum(value < 0.95 for value in baseline_v)
opt_hours_low = sum(value < 0.95 for value in optimized_v)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = "D5DAE1", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_shading(paragraph, fill: str, border_color: str = BLUE) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)
    borders.append(left)
    p_pr.append(borders)
    ind = paragraph.paragraph_format
    ind.left_indent = Inches(0.12)
    ind.right_indent = Inches(0.08)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, size=9, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(
    doc: Document,
    text: str = "",
    *,
    bold_prefix: str | None = None,
    italic: bool = False,
    after: float | None = None,
):
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        second = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(second, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, italic=italic)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    set_run_font(paragraph.add_run(text))
    return paragraph


def start_numbered_list(doc: Document) -> int:
    """Create a restartable instance of Word's native List Number definition."""
    numbering = doc.part.numbering_part.element
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering
        if node.tag == qn("w:num")
    ]
    num_id = max(num_ids, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), "7")
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def add_numbered(doc: Document, text: str, num_id: int):
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.append(num_pr)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_callout(
    doc: Document,
    label: str,
    text: str,
    *,
    fill: str = LIGHT_BLUE,
    border: str = BLUE,
):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.12
    set_paragraph_shading(paragraph, fill, border)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=NAVY)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, color=BLACK)
    return paragraph


def add_source_note(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=8.5, color=MUTED, italic=True)
    return paragraph


def add_figure(
    doc: Document,
    image_name: str,
    caption: str,
    *,
    width: float = 6.25,
    alt_text: str,
):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    inline = paragraph.add_run().add_picture(str(RESULTS / image_name), width=Inches(width))
    inline._inline.docPr.set("descr", alt_text)
    caption_p = doc.add_paragraph(style="Caption")
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(2)
    caption_p.paragraph_format.space_after = Pt(8)
    caption_p.paragraph_format.keep_together = True
    run = caption_p.add_run(caption)
    set_run_font(run, size=9, color=MUTED, italic=True)
    return paragraph, caption_p


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    *,
    font_size: float = 9.2,
    first_col_bold: bool = True,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.allow_autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=font_size, color=NAVY, bold=True)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            set_cell_borders(cell)
            if row_index % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if col_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(str(value))
            set_run_font(
                run,
                size=font_size,
                color=BLACK,
                bold=bool(first_col_bold and col_index == 0),
            )
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 90, "bottom": 90, "start": 120, "end": 120},
    )
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

caption_style = styles["Caption"]
caption_style.font.name = "Calibri"
caption_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
caption_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
caption_style.font.size = Pt(9)
caption_style.font.italic = True
caption_style.font.color.rgb = RGBColor.from_string(MUTED)

for list_style_name in ("List Bullet", "List Number"):
    style = styles[list_style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(11)

header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_p.paragraph_format.space_after = Pt(0)
header_run = header_p.add_run("DUCK-CURVE OPTIMIZATION | TECHNICAL RESULTS REPORT")
set_run_font(header_run, size=8.5, color=MUTED, bold=True)

footer_p = section.footer.paragraphs[0]
add_page_number(footer_p)
first_footer_p = section.first_page_footer.paragraphs[0]
first_footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(
    first_footer_p.add_run("Prepared from verified project-generated outputs | July 2026"),
    size=8.5,
    color=MUTED,
)

doc.core_properties.title = (
    "Technical Results Report: Network-Aware PV-BESS Optimization for Duck-Curve Flattening"
)
doc.core_properties.subject = "IEEE 33-bus PV-BESS optimization and resilience results"
doc.core_properties.author = ""
doc.core_properties.keywords = "duck curve, BESS, PV, IEEE 33-bus, resilience, SSS"

# Cover page: editorial_cover pattern with a restrained technical-report override.
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(74)
kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_after = Pt(14)
set_run_font(
    kicker.add_run("TECHNICAL RESULTS REPORT"),
    size=10.5,
    color=GOLD,
    bold=True,
)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(10)
set_run_font(
    title.add_run("Network-Aware PV-BESS Optimization"),
    size=29,
    color=NAVY,
    bold=True,
)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(4)
set_run_font(
    subtitle.add_run("Duck-Curve Flattening, Voltage Performance, and Resilience"),
    size=15,
    color=DARK_BLUE,
)
case_p = doc.add_paragraph()
case_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
case_p.paragraph_format.space_after = Pt(34)
set_run_font(
    case_p.add_run("IEEE 33-bus radial feeder | 24-hour study | 3 PV + 3 BESS"),
    size=10.5,
    color=MUTED,
    italic=True,
)

metric = doc.add_paragraph()
metric.alignment = WD_ALIGN_PARAGRAPH.CENTER
metric.paragraph_format.space_before = Pt(8)
metric.paragraph_format.space_after = Pt(8)
set_paragraph_shading(metric, LIGHT_BLUE, BLUE)
set_run_font(metric.add_run("VERIFIED SEED-42 RESULT\n"), size=10, color=BLUE, bold=True)
set_run_font(
    metric.add_run(f"{f(improvement, 'sum_slope_sq_kW2'):.2f}% SSS reduction"),
    size=22,
    color=NAVY,
    bold=True,
)

status = doc.add_paragraph()
status.alignment = WD_ALIGN_PARAGRAPH.CENTER
status.paragraph_format.space_before = Pt(26)
status.paragraph_format.space_after = Pt(3)
set_run_font(status.add_run("Prepared for academic review"), size=11, color=BLACK, bold=True)
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(date_p.add_run("Results version: July 24, 2026"), size=10, color=MUTED)
scope = doc.add_paragraph()
scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
scope.paragraph_format.space_before = Pt(40)
set_run_font(
    scope.add_run(
        "Evidence status: complete engineering case study; multi-seed statistical validation pending."
    ),
    size=9.5,
    color=RED,
    italic=True,
)

doc.add_page_break()

add_heading(doc, "Executive Summary", 1)
add_body(
    doc,
    (
        "The verified seed-42 solution reduces the sum of squared hourly net-load slopes "
        f"(SSS) from {f(baseline, 'sum_slope_sq_kW2'):,.2f} to "
        f"{f(optimized, 'sum_slope_sq_kW2'):,.2f} kW², an improvement of "
        f"{f(improvement, 'sum_slope_sq_kW2'):.2f}%. The optimized placement is PV at "
        f"buses {optimized['pv_buses'].replace(';', ', ')} and BESS at buses "
        f"{optimized['bess_buses'].replace(';', ', ')}."
    ),
)
add_callout(
    doc,
    "Primary conclusion",
    (
        "The model achieves the requested approximately 81% duck-curve reduction while "
        "maintaining per-unit BESS power, SOC bounds, and exact daily cycle closure. "
        "The differentiated SOC trajectories arise from objective-accepted network-aware "
        "redispatch, not from an imposed participation pattern."
    ),
)

add_body(doc, "Table 1. Headline before-and-after performance", italic=True, after=4)
add_table(
    doc,
    ["Metric", "Before", "After", "Change"],
    [
        [
            "SSS (kW²)",
            f"{f(baseline, 'sum_slope_sq_kW2'):,.2f}",
            f"{f(optimized, 'sum_slope_sq_kW2'):,.2f}",
            f"{f(improvement, 'sum_slope_sq_kW2'):.2f}% lower",
        ],
        [
            "Peak net load (kW)",
            f"{f(baseline, 'peak_kW'):,.2f}",
            f"{f(optimized, 'peak_kW'):,.2f}",
            f"{f(improvement, 'peak_kW'):.2f}% lower",
        ],
        [
            "Net-load range (kW)",
            f"{f(baseline, 'range_kW'):,.2f}",
            f"{f(optimized, 'range_kW'):,.2f}",
            f"{f(improvement, 'range_kW'):.2f}% lower",
        ],
        [
            "Peak hourly ramp (kW/h)",
            f"{f(baseline, 'ramp_peak_kWh'):,.2f}",
            f"{f(optimized, 'ramp_peak_kWh'):,.2f}",
            f"{f(improvement, 'ramp_peak_kWh'):.2f}% lower",
        ],
        [
            "Expected ENS (kWh/outage)",
            f"{f(res_before, 'eens_kwh'):,.2f}",
            f"{f(res_after, 'eens_kwh'):,.2f}",
            f"{pct_lower(f(res_before, 'eens_kwh'), f(res_after, 'eens_kwh')):.2f}% lower",
        ],
        [
            "Mean resilience index",
            f"{f(res_before, 'resilience_index'):.4f}",
            f"{f(res_after, 'resilience_index'):.4f}",
            f"{pct_higher(f(res_before, 'resilience_index'), f(res_after, 'resilience_index')):.2f}% higher",
        ],
    ],
    [3100, 1900, 1900, 2460],
)
add_source_note(doc, "Source: results_summary.csv and resilience_indices.csv (seed 42).")

add_heading(doc, "Key Findings", 2)
add_bullet(
    doc,
    (
        f"The midday negative trough is removed: net load changes from "
        f"{f(baseline, 'trough_kW'):,.2f} kW to {f(optimized, 'trough_kW'):,.2f} kW."
    ),
)
add_bullet(
    doc,
    (
        f"All BESS units remain within 0.6–2.7 MWh, use less than 0.60 MW of the "
        f"1.0 MW unit rating, and close the daily SOC cycle with zero reported error."
    ),
)
add_bullet(
    doc,
    (
        f"The worst hourly feeder voltage improves from {worst_base_v:.5f} to "
        f"{worst_opt_v:.5f} p.u.; however, the optimized value remains "
        f"{0.95 - worst_opt_v:.5f} p.u. below the 0.95 threshold."
    ),
)
add_bullet(
    doc,
    (
        f"Four-hour outage EENS falls by "
        f"{pct_lower(f(res_before, 'eens_kwh'), f(res_after, 'eens_kwh')):.2f}%, "
        f"and zero-ENS scenarios increase from {scenario_stats['before']['zero_ens']} "
        f"to {scenario_stats['after']['zero_ens']} of 768."
    ),
)
add_callout(
    doc,
    "Evidence limitation",
    (
        "The reported optimum is from one random seed. It is valid as the verified "
        "case result, but it should not be presented as statistically robust until a "
        "multi-seed experiment is completed."
    ),
    fill=LIGHT_GOLD,
    border=GOLD,
)

doc.add_page_break()

add_heading(doc, "1. Study Scope and System Model", 1)
add_body(
    doc,
    (
        "The study evaluates coordinated placement and hourly dispatch of photovoltaic "
        "(PV) generation and battery energy storage systems (BESS) on the IEEE 33-bus "
        "radial distribution feeder. The analysis uses a deterministic 24-hour "
        "residential-load profile and solar-production profile."
    ),
)
add_heading(doc, "1.1 Model Configuration", 2)
for item in (
    "Three 1.0 MW PV units; total installed PV capacity is 3.0 MW.",
    "Three BESS units; total power and energy ratings are 3.0 MW and 9.0 MWh.",
    "Each BESS has a 1.0 MW power limit and 3.0 MWh energy capacity.",
    "Charge and discharge efficiencies are each 95%.",
    "SOC is constrained to 20%–90% of rated energy (0.6–2.7 MWh per unit).",
    "Each BESS must return to its own initial SOC at the end of hour 24.",
    "Voltage magnitude limits are modeled as 0.95–1.05 p.u.",
):
    add_bullet(doc, item)

add_heading(doc, "1.2 Compared Operating Cases", 2)
case_list_id = start_numbered_list(doc)
add_numbered(
    doc,
    "No-PV/no-BESS reference case for the unmodified feeder load.",
    case_list_id,
)
add_numbered(
    doc,
    "PV-only duck-curve baseline with the original baseline PV placement and no BESS support.",
    case_list_id,
)
add_numbered(
    doc,
    "Optimized PV+BESS case with jointly selected buses and hourly battery schedules.",
    case_list_id,
)
add_figure(
    doc,
    "fig2_load_and_pv.png",
    "Figure 1. Normalized 24-hour load and PV profiles used by the model.",
    width=4.8,
    alt_text="Normalized residential load and photovoltaic generation profiles over 24 hours.",
)
optimization_heading = add_heading(doc, "2. Optimization Formulation", 1)
optimization_heading.paragraph_format.page_break_before = True
add_heading(doc, "2.1 Objectives", 2)
add_body(
    doc,
    (
        "The primary objective is the sum of squared consecutive hourly net-load slopes:"
    ),
)
equation = doc.add_paragraph()
equation.alignment = WD_ALIGN_PARAGRAPH.CENTER
equation.paragraph_format.space_before = Pt(6)
equation.paragraph_format.space_after = Pt(10)
set_run_font(
    equation.add_run("SSS = sum from t=0 to 22 of (Pnet,t+1 - Pnet,t)²"),
    size=12,
    color=NAVY,
    italic=True,
)
add_body(
    doc,
    (
        "A lower SSS corresponds to a smoother net-load trajectory. The second objective "
        "is the SALEDI outage surrogate. Feeder losses, voltage excursions, duplicate "
        "placements, SOC bounds, unit power limits, and cycle closure are evaluated in "
        "the candidate fitness calculation."
    ),
)
add_heading(doc, "2.2 Search Procedure", 2)
search_list_id = start_numbered_list(doc)
for item in (
    "Construct an efficiency-aware analytical load-leveling schedule as a feasible warm start.",
    "Apply MO-EZOA search to placement and 24-hour dispatch variables.",
    "Polish PV and BESS bus placements by constrained discrete coordinate search.",
    "Refine aggregate BESS dispatch while projecting to power, SOC-span, and cycle constraints.",
    "Refine each physical BESS independently.",
    (
        "Apply symmetric pairwise redispatch: transfer power between two units at one hour "
        "and reverse the transfer at another hour in the same operating regime. This "
        "preserves aggregate power and each unit’s efficiency-aware daily energy balance."
    ),
):
    add_numbered(doc, item, search_list_id)
add_callout(
    doc,
    "Anti-bias safeguard",
    (
        "No target SSS, target reduction percentage, SOC separation reward, participation "
        "factor, or final bus placement is encoded. Pairwise redispatch candidates are "
        "symmetric and are retained only when the evaluated objectives improve."
    ),
)
add_heading(doc, "2.3 Verified Run Settings", 2)
add_table(
    doc,
    ["Parameter", "Verified value", "Interpretation"],
    [
        ["Random seed", "42", "Current reported case"],
        ["Population", "6", "Quick verified configuration"],
        ["MO-EZOA iterations", "4", "Quick verified configuration"],
        ["Elite polish trials", "20", "Constrained derivative-free refinement"],
        ["PV/BESS counts", "3 / 3", "Fixed resource counts"],
        ["Losses in net load", "Enabled", "Placement affects primary objective"],
    ],
    [2400, 1700, 5260],
)
add_source_note(doc, "Source: configs/verified_81pct.json.")

doc.add_page_break()

add_heading(doc, "3. Duck-Curve and Ramp Results", 1)
add_figure(
    doc,
    "fig4_netload_before_after.png",
    "Figure 2. Net-load profiles before and after optimization.",
    width=6.25,
    alt_text="Hourly no-DER, PV-only baseline, and optimized PV plus BESS net-load curves.",
)
add_body(
    doc,
    (
        "The optimized schedule raises the PV-driven midday trough and reduces the evening "
        "peak. Relative to the PV-only baseline, the net-load range decreases by "
        f"{f(improvement, 'range_kW'):.2f}% and the peak hourly ramp decreases by "
        f"{f(improvement, 'ramp_peak_kWh'):.2f}%. The remaining 2.17 MW range reflects "
        "the finite 3 MW/9 MWh BESS resource and cycle-closure requirement."
    ),
)
add_body(doc, "Table 2. Detailed duck-curve performance", italic=True, after=4)
add_table(
    doc,
    ["Metric", "PV-only baseline", "Optimized", "Improvement"],
    [
        [
            "SSS (kW²)",
            f"{f(baseline, 'sum_slope_sq_kW2'):,.2f}",
            f"{f(optimized, 'sum_slope_sq_kW2'):,.2f}",
            f"{f(improvement, 'sum_slope_sq_kW2'):.2f}%",
        ],
        [
            "Variance (kW²)",
            f"{f(baseline, 'variance_kW2'):,.2f}",
            f"{f(optimized, 'variance_kW2'):,.2f}",
            f"{f(improvement, 'variance_kW2'):.2f}%",
        ],
        [
            "Peak (kW)",
            f"{f(baseline, 'peak_kW'):,.2f}",
            f"{f(optimized, 'peak_kW'):,.2f}",
            f"{f(improvement, 'peak_kW'):.2f}%",
        ],
        [
            "Trough (kW)",
            f"{f(baseline, 'trough_kW'):,.2f}",
            f"{f(optimized, 'trough_kW'):,.2f}",
            "Negative export removed",
        ],
        [
            "Range (kW)",
            f"{f(baseline, 'range_kW'):,.2f}",
            f"{f(optimized, 'range_kW'):,.2f}",
            f"{f(improvement, 'range_kW'):.2f}%",
        ],
        [
            "Mean ramp (kW/h)",
            f"{f(baseline, 'ramp_mean_kWh'):,.2f}",
            f"{f(optimized, 'ramp_mean_kWh'):,.2f}",
            f"{f(improvement, 'ramp_mean_kWh'):.2f}%",
        ],
        [
            "Peak ramp (kW/h)",
            f"{f(baseline, 'ramp_peak_kWh'):,.2f}",
            f"{f(optimized, 'ramp_peak_kWh'):,.2f}",
            f"{f(improvement, 'ramp_peak_kWh'):.2f}%",
        ],
    ],
    [2700, 2200, 2000, 2460],
)
add_source_note(doc, "Source: results_summary.csv.")

doc.add_page_break()

add_heading(doc, "4. BESS Dispatch and SOC Verification", 1)
add_figure(
    doc,
    "fig5_soc_differentiated.png",
    "Figure 3. Actual independently optimized SOC trajectories for all BESS units.",
    width=6.25,
    alt_text=(
        "State-of-charge trajectories for BESS at buses 28, 15, and 32 with minimum "
        "and maximum energy limits."
    ),
)
add_body(
    doc,
    (
        f"The maximum observed pairwise SOC separation is "
        f"{max(float(row['max_soc_separation_kwh']) for row in audit_rows):.1f} kWh. "
        "The separation is a physical optimization outcome: network-aware paired "
        "redispatch preserves total feeder-level BESS power while changing which unit "
        "provides that power at different hours."
    ),
)
add_body(doc, "Table 3. Per-unit physical constraint audit", italic=True, after=4)
audit_table_rows = []
for row in audit_rows:
    audit_table_rows.append(
        [
            f"BESS {row['unit']}\nBus {row['bus']}",
            f"{float(row['max_abs_power_mw']):.3f} / 1.000",
            (
                f"{float(row['soc_min_observed_mwh']):.3f}–"
                f"{float(row['soc_max_observed_mwh']):.3f}"
            ),
            (
                f"{float(row['initial_soc_mwh']):.3f} / "
                f"{float(row['final_soc_mwh']):.3f}"
            ),
            (
                f"{float(row['cycle_error_kwh']):.3f}\n"
                f"{float(row['max_soc_separation_kwh']):.1f}"
            ),
        ]
    )
add_table(
    doc,
    [
        "Unit / bus",
        "Max power\nused / limit (MW)",
        "Observed SOC\nrange (MWh)",
        "Initial / final\nSOC (MWh)",
        "Cycle error /\nmax separation (kWh)",
    ],
    audit_table_rows,
    [1500, 1900, 1900, 1900, 2160],
    font_size=8.7,
)
add_source_note(doc, "Source: bess_constraint_audit.csv.")
add_callout(
    doc,
    "Constraint finding",
    (
        "Every unit satisfies the 1 MW power rating and the 0.6–2.7 MWh SOC window. "
        "Reported cycle error is zero for all units; therefore, the differentiated "
        "curves are not produced by unbalanced end-of-day energy."
    ),
)

doc.add_page_break()

add_heading(doc, "5. Voltage Performance", 1)
add_figure(
    doc,
    "fig6_voltage_profile.png",
    "Figure 4. Literature-style bus-voltage profile at the DER-active critical hour.",
    width=6.2,
    alt_text=(
        "Bus voltage profiles for no-DER, PV-only, and optimized PV plus BESS cases "
        "with 0.95 and 1.05 per-unit limits."
    ),
)
add_body(doc, "Table 4. Minimum-voltage assessment", italic=True, after=4)
add_table(
    doc,
    ["Indicator", "PV-only baseline", "Optimized", "Assessment"],
    [
        [
            "Worst hourly minimum",
            f"{worst_base_v:.5f} p.u.",
            f"{worst_opt_v:.5f} p.u.",
            f"+{worst_opt_v - worst_base_v:.5f} p.u.",
        ],
        [
            "Hours below 0.95 p.u.",
            str(base_hours_low),
            str(opt_hours_low),
            f"{base_hours_low - opt_hours_low} fewer hours",
        ],
        [
            "Worst optimized hour",
            "—",
            f"Hour {worst_opt_hour}",
            f"{0.95 - worst_opt_v:.5f} p.u. below limit",
        ],
    ],
    [2500, 1900, 1900, 3060],
)
add_source_note(doc, "Source: results_hourly.csv.")
add_callout(
    doc,
    "Voltage compliance limitation",
    (
        f"The optimized worst voltage is {worst_opt_v:.5f} p.u. at hour "
        f"{worst_opt_hour}. It is improved substantially but remains "
        f"{0.95 - worst_opt_v:.5f} p.u. below 0.95. The current result should be "
        "described as near-compliant, not fully compliant."
    ),
    fill=LIGHT_GOLD,
    border=GOLD,
)

doc.add_page_break()

add_heading(doc, "5.1 Complete 24-Hour Voltage Assessment", 2)
add_figure(
    doc,
    "fig6b_voltage_complete_24h.png",
    "Figure 5. Complete spatial-temporal voltage assessment for all buses and hours.",
    width=6.25,
    alt_text=(
        "Daily minimum voltage by bus, hourly feeder minimum voltage, and full "
        "hour-by-bus voltage heat maps before and after optimization."
    ),
)
add_body(
    doc,
    (
        "The heat maps show that the optimization improves the evening undervoltage "
        "region across the feeder. The remaining small violation is concentrated around "
        "the evening peak rather than representing broad all-day noncompliance."
    ),
)
add_body(
    doc,
    (
        "The voltage calculation used by the quick verified case is the analytical "
        "DistFlow approximation. A full AC or backward-forward-sweep power-flow check "
        "should be included before making a final network-compliance claim."
    ),
)

doc.add_page_break()

add_heading(doc, "6. Resilience Assessment", 1)
add_body(
    doc,
    (
        "Resilience is evaluated for every feeder-line outage and every possible outage "
        "start hour. With 32 lines and 24 start hours, each before/after case contains "
        "768 four-hour islanding scenarios. Local PV serves downstream island load first; "
        "BESS support is limited by outage-start SOC, power rating, and usable energy."
    ),
)
add_figure(
    doc,
    "fig9_resilience_indices.png",
    "Figure 6. Before-and-after resilience indicators for four-hour line outages.",
    width=6.15,
    alt_text=(
        "Comparison of expected and worst-case energy not served, load served, "
        "resilience index, and SALEDI before and after optimization."
    ),
)
add_body(doc, "Table 5. Aggregate resilience results", italic=True, after=4)
add_table(
    doc,
    ["Metric", "Before", "After", "Change"],
    [
        [
            "Expected ENS (kWh)",
            f"{f(res_before, 'eens_kwh'):,.2f}",
            f"{f(res_after, 'eens_kwh'):,.2f}",
            f"{pct_lower(f(res_before, 'eens_kwh'), f(res_after, 'eens_kwh')):.2f}% lower",
        ],
        [
            "Worst-case ENS (kWh)",
            f"{f(res_before, 'worst_case_ens_kwh'):,.2f}",
            f"{f(res_after, 'worst_case_ens_kwh'):,.2f}",
            f"{pct_lower(f(res_before, 'worst_case_ens_kwh'), f(res_after, 'worst_case_ens_kwh')):.2f}% lower",
        ],
        [
            "Aggregate load served",
            f"{f(res_before, 'load_served_percent'):.2f}%",
            f"{f(res_after, 'load_served_percent'):.2f}%",
            f"+{f(res_after, 'load_served_percent') - f(res_before, 'load_served_percent'):.2f} points",
        ],
        [
            "Mean resilience index",
            f"{f(res_before, 'resilience_index'):.4f}",
            f"{f(res_after, 'resilience_index'):.4f}",
            f"{pct_higher(f(res_before, 'resilience_index'), f(res_after, 'resilience_index')):.2f}% higher",
        ],
        [
            "SALEDI",
            f"{f(res_before, 'saledi'):.3f}",
            f"{f(res_after, 'saledi'):.3f}",
            f"{pct_lower(f(res_before, 'saledi'), f(res_after, 'saledi')):.2f}% lower",
        ],
    ],
    [2900, 1700, 1700, 3060],
)
add_source_note(doc, "Source: resilience_indices.csv.")

doc.add_page_break()

add_heading(doc, "6.1 Scenario Distributions", 2)
add_figure(
    doc,
    "fig9b_resilience_violins.png",
    "Figure 7. Violin distributions across all 768 before and 768 after outage scenarios.",
    width=6.25,
    alt_text=(
        "Violin plots of scenario energy not served and load served percentages before "
        "and after optimization."
    ),
)
add_body(doc, "Table 6. Distribution-level resilience statistics", italic=True, after=4)
add_table(
    doc,
    ["Statistic", "Before", "After", "Interpretation"],
    [
        [
            "Mean scenario ENS (kWh)",
            f"{scenario_stats['before']['ens_mean']:,.1f}",
            f"{scenario_stats['after']['ens_mean']:,.1f}",
            "Lower after optimization",
        ],
        [
            "Median scenario ENS (kWh)",
            f"{scenario_stats['before']['ens_median']:,.1f}",
            f"{scenario_stats['after']['ens_median']:,.1f}",
            f"{pct_lower(scenario_stats['before']['ens_median'], scenario_stats['after']['ens_median']):.1f}% lower",
        ],
        [
            "Mean scenario load served",
            f"{scenario_stats['before']['served_mean']:.1f}%",
            f"{scenario_stats['after']['served_mean']:.1f}%",
            f"+{scenario_stats['after']['served_mean'] - scenario_stats['before']['served_mean']:.1f} points",
        ],
        [
            "Zero-ENS scenarios",
            f"{scenario_stats['before']['zero_ens']} / 768",
            f"{scenario_stats['after']['zero_ens']} / 768",
            f"+{scenario_stats['after']['zero_ens'] - scenario_stats['before']['zero_ens']} scenarios",
        ],
    ],
    [2700, 1700, 1700, 3260],
)
add_source_note(doc, "Source: resilience_scenarios.csv.")
add_body(
    doc,
    (
        "The distribution remains broad because outage impact depends strongly on the "
        "outaged line and start hour. The after-optimization violin nevertheless shifts "
        "toward lower unserved energy and higher served load, indicating that the aggregate "
        "improvement is not caused by only a few favorable scenarios."
    ),
)

doc.add_page_break()

add_heading(doc, "6.2 Line-Specific Resilience Effects", 2)
add_figure(
    doc,
    "fig10_line_outage_resilience.png",
    "Figure 8. Mean four-hour unserved energy by outaged IEEE 33-bus feeder line.",
    width=6.25,
    alt_text=(
        "Line-by-line expected energy not served before and after optimization, with "
        "areas of improvement highlighted."
    ),
)
add_body(
    doc,
    (
        "Benefits are not uniform across every line. The optimized placement improves "
        "the aggregate resilience indices, but some local outage locations show smaller "
        "benefits or adverse trade-offs. This is an important engineering result rather "
        "than a plotting defect: a placement selected for system-level SSS, voltage, and "
        "outage objectives does not guarantee improvement for every individual island."
    ),
)
add_callout(
    doc,
    "Recommended follow-up",
    (
        "If line-specific equity is a design requirement, add a worst-line EENS or "
        "maximum-regret constraint instead of filtering unfavorable line results after "
        "optimization."
    ),
)

doc.add_page_break()

add_heading(doc, "7. Reproducibility and Multi-Seed Validation", 1)
add_callout(
    doc,
    "Answer",
    (
        "Yes. Different seeds should be run before the 82.84% result is presented as "
        "robust. Seed 42 is a verified feasible result, but a single stochastic run "
        "cannot quantify solution variability or placement stability."
    ),
)
add_body(
    doc,
    (
        "The analytical warm start and deterministic constrained polish reduce randomness, "
        "but MO-EZOA initialization and search moves remain seed-dependent. The project’s "
        "paper-scale configuration also records that earlier single-seed outcomes varied "
        "substantially, reinforcing the need for a formal seed study."
    ),
)
add_heading(doc, "7.1 Recommended Validation Plan", 2)
add_body(doc, "Table 7. Staged multi-seed protocol", italic=True, after=4)
add_table(
    doc,
    ["Stage", "Seeds / budget", "Purpose", "Decision rule"],
    [
        [
            "Screening",
            "10 seeds\n42, 1042, …, 9042",
            "Check reproducibility of the verified quick case",
            "Continue if ≥9/10 achieve ≥81% and SSS CV ≤1%",
        ],
        [
            "Robustness",
            "20–30 seeds",
            "Estimate mean, spread, worst case, and placement frequency",
            "Required if screening variability exceeds thresholds",
        ],
        [
            "Paper-scale confirmation",
            "≥10 seeds with larger population/iterations",
            "Confirm that conclusions are not artifacts of quick settings",
            "Report all seeds; do not select only the best",
        ],
    ],
    [1500, 1900, 2850, 3110],
    font_size=8.6,
)
add_source_note(
    doc,
    (
        "The ≥9/10 and coefficient-of-variation thresholds are recommended project "
        "acceptance criteria, not universal statistical standards."
    ),
)

add_heading(doc, "7.2 Statistics to Retain for Every Seed", 2)
for item in (
    "Best feasible SSS and percentage reduction relative to the identical PV-only baseline.",
    "Mean, median, standard deviation, coefficient of variation, minimum, and maximum SSS reduction.",
    "Success rate for achieving at least 81% reduction.",
    "PV and BESS bus-placement frequency and the number of distinct placement patterns.",
    "Minimum/maximum voltage, number of hours below 0.95 p.u., and all physical constraint residuals.",
    "SALEDI, EENS, worst-case ENS, resilience index, and runtime.",
):
    add_bullet(doc, item)

add_heading(doc, "7.3 Suggested Commands", 2)
add_body(
    doc,
    (
        "Copy configs/verified_81pct.json to configs/seed_validation_10.json, change "
        "\"n_seeds\" from 1 to 10, and run:"
    ),
)
code_p = doc.add_paragraph()
code_p.paragraph_format.left_indent = Inches(0.25)
code_p.paragraph_format.right_indent = Inches(0.25)
code_p.paragraph_format.space_before = Pt(4)
code_p.paragraph_format.space_after = Pt(10)
set_paragraph_shading(code_p, "F7F7F7", MUTED)
set_run_font(
    code_p.add_run(
        "python run_paper_case.py --config configs/seed_validation_10.json "
        "--output-dir results_seed_validation"
    ),
    name="Consolas",
    size=9,
    color=BLACK,
)
add_body(
    doc,
    (
        "The present multi-seed runner merges the Pareto fronts and prints per-seed best "
        "SSS values. For manuscript-quality evidence, retain a CSV row for every seed "
        "rather than relying only on the merged best solution."
    ),
)

add_heading(doc, "8. Limitations and Required Follow-Up", 1)
limitations = (
    (
        "Single-seed evidence",
        "The current 82.84% result is seed 42 only. Multi-seed statistics are pending.",
    ),
    (
        "Quick search budget",
        "Population 6 and four MO-EZOA iterations are intentionally small; deterministic polishing performs much of the final improvement.",
    ),
    (
        "Residual voltage violation",
        f"The worst optimized voltage is {worst_opt_v:.5f} p.u., marginally below 0.95 p.u.",
    ),
    (
        "Power-flow fidelity",
        "The verified run uses analytical DistFlow voltage/loss calculations; confirm the final case with a full AC or backward-forward-sweep model.",
    ),
    (
        "Deterministic profiles",
        "The study uses one load day and one PV day. Forecast error, weather variability, and seasonal cases are not represented.",
    ),
    (
        "Resilience abstraction",
        "The four-hour islanding model omits protection, switching time, restoration sequencing, cold-load pickup, and PV-surplus charging during the outage.",
    ),
    (
        "Economic scope",
        "Battery degradation, investment cost, market revenue, and lifecycle economics are not objective terms in the verified case.",
    ),
)
for label, text in limitations:
    add_body(doc, f"{label}: {text}", bold_prefix=f"{label}:")

add_heading(doc, "Recommended Next Actions", 2)
action_list_id = start_numbered_list(doc)
for action in (
    "Run the 10-seed screening study and save one complete result row per seed.",
    "If screening is variable, expand to 20–30 seeds and increase population/iterations.",
    "Verify the selected placements and dispatch with a full power-flow solver.",
    "Add a stricter voltage constraint or voltage-focused repair if full 0.95 p.u. compliance is mandatory.",
    "Evaluate multiple seasonal load/PV days and uncertainty scenarios.",
    "Retain line-level resilience trade-offs instead of reporting only aggregate improvements.",
):
    add_numbered(doc, action, action_list_id)

doc.add_page_break()

add_heading(doc, "9. Conclusions", 1)
add_body(
    doc,
    (
        f"The seed-42 solution demonstrates that coordinated placement and dispatch of "
        f"3 MW/9 MWh of BESS with 3 MW of PV can reduce the PV-only duck-curve SSS by "
        f"{f(improvement, 'sum_slope_sq_kW2'):.2f}% on the modeled IEEE 33-bus feeder. "
        f"Peak net load falls by {f(improvement, 'peak_kW'):.2f}%, peak ramp falls by "
        f"{f(improvement, 'ramp_peak_kWh'):.2f}%, and the PV-induced negative midday "
        "trough is removed."
    ),
)
add_body(
    doc,
    (
        f"The BESS schedules are physically auditable and differentiated, with a maximum "
        f"SOC separation of {max(float(row['max_soc_separation_kwh']) for row in audit_rows):.1f} "
        "kWh, no unit power violation, no SOC-bound violation, and zero reported cycle "
        "error. Resilience also improves: EENS decreases by "
        f"{pct_lower(f(res_before, 'eens_kwh'), f(res_after, 'eens_kwh')):.2f}% and the "
        f"mean resilience index increases by "
        f"{pct_higher(f(res_before, 'resilience_index'), f(res_after, 'resilience_index')):.2f}%."
    ),
)
add_body(
    doc,
    (
        "Two qualifications remain essential. First, minimum voltage is near but not fully "
        "inside the 0.95 p.u. limit. Second, the result is based on one seed and should be "
        "described as the best verified case until a multi-seed robustness study is complete."
    ),
)
add_callout(
    doc,
    "Recommended claim wording",
    (
        "“The verified seed-42 case achieved an 82.84% SSS reduction under the modeled "
        "constraints. Multi-seed and full power-flow validation are recommended before "
        "generalizing this result.”"
    ),
)

doc.add_page_break()
add_heading(doc, "Appendix A. Result Files Used", 1)
for name in (
    "results_summary.csv",
    "results_hourly.csv",
    "bess_constraint_audit.csv",
    "resilience_indices.csv",
    "resilience_scenarios.csv",
    "fig4_netload_before_after.png",
    "fig5_soc_differentiated.png",
    "fig6_voltage_profile.png",
    "fig6b_voltage_complete_24h.png",
    "fig9_resilience_indices.png",
    "fig9b_resilience_violins.png",
    "fig10_line_outage_resilience.png",
):
    add_bullet(doc, f"results_81pct/{name}")

REPORTS.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
